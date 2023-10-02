import datetime
import os
import sys
import shutil
import ntpath

from docx import Document


class TerminalColors:
    HEADER = '\033[95m'
    OK_BLUE = '\033[94m'
    OK_CYAN = '\033[96m'
    OK_GREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    END_COLOR = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    RESET = '\x1b[0m'


def loading(curr_item: int, all_items: int, debug=False):
    if debug:
        return
    percent = ((curr_item + 1) / all_items) * 100
    os.system("cls")
    print("|", end='')
    for j in range(0, 101):
        if j > percent:
            print('_', end='')
        else:
            print('=', end='')
    print(f'> {percent} ({curr_item + 1} out of {all_items})')


quotes = {'”', '"'}
special_chars = {'?', '!', ',', ':', '.'}
incorrect_patterns = [
    # These useless chars
    ['»', ''],
    ['«', ''],
    [' ', ' '],

    # Multiple punctuation
    ['  ', ' '], # Special space
    ['  ', ' '], # Special space
    # ['. .', '.'],
    ['!.', '!'],
    ['? !', '?!'],
    ['! ?', '!?'],
    ['.!', '.'],
    [' :', ':'],

    # Bad quotes
    ['’’', '”'],  # Upper
    [',,', '„'],  # Lower
    [', ,', ','],
    ['‘’', ','],

    [', és', ' és'],
    [',és', ' és'],
    ['úgy ', 'úgy, '],
    ['elöre', 'előre'],

    # Random words
    ['CIME', 'CÍME'],
    ['cime', 'címe'],
    ['dicsőités', 'dicsőítés'],
    ['Dicsőités', 'Dicsőítés'],
    ['dicsőit', 'dicsőít'],
    ['Dicsőit', 'Dicsőít'],
    ['íge', 'ige'],
    ['ídéz', 'idéz'],
    ['Hungarian Bible Easy-to-read Version', 'Biblia-Egyszerű fordítás'],
    ['Hungarian Bible Easy-to-Read Version', 'Biblia-Egyszerű fordítás'],
    ['Magyar Biblia: Egyszerű forditás', 'Biblia-Egyszerű fordítás'],
    ['Magyar Biblia: Egyszerű fordítás', 'Biblia-Egyszerű fordítás'],

    # Bible books
    ['Korinthus', 'Korintus'],
    ['Cselekedetek', 'ApCsel'],
    ['Kolosséiakhoz', 'Kolossé'],

    # Sites
    ['www. ', 'www.'],
    ['. org', '.org'],
    ['http: //', 'http://'],
    ['https: //', 'https://']
]


def my_print(my_text: str, end, debug=False):
    if debug:
        print(my_text, end=end)


def remove_patterns(paragraph, debug=False) -> str:
    p = paragraph._p
    runs = paragraph.runs
    for index, run in enumerate(runs):
        if debug:
            print(run.text)

        for i in range(8, 1, -1):
            replace = " " * i
            run.text = run.text.replace(replace, " ")
            if index + 1 < len(runs) and index - 1 >= 0:
                next_run = runs[index + 1]
                prev_run = runs[index - 1]
                if len(next_run.text) > 0:
                    combined_text = run.text + next_run.text
                    if replace in combined_text:
                        if len(run.text) > 0 and run.text[len(run.text) - 1] == " " and next_run.text[0] == " ":
                            next_run.text = next_run.text[1:]
                        if len(run.text) > 0 and run.text[0] == " " \
                                and len(prev_run.text) > 0 \
                                and prev_run.text[len(prev_run.text) - 1] == " ":
                            prev_run.text = prev_run.text[:len(prev_run.text) - 1]
                    if len(run.text) > 0 and (run.text[len(run.text) - 1] in special_chars) \
                            and next_run.text[0] != " " \
                            and (next_run.text[0] not in quotes):
                        next_run.text = " " + next_run.text
                        if debug:
                            print(next_run)
        limit = len(run.text)
        for i in range(0, limit):
            if (i - 1 >= 0) and i + 1 < len(run.text):
                next_char = run.text[i + 1]
                prev_char = run.text[i - 1]
                if (run.text[i] in special_chars) \
                        and (next_char != ' ') \
                        and (next_char not in quotes and not(prev_char.isnumeric() and next_char.isnumeric())):
                    run.text = run.text[:i + 1] + " " + run.text[i + 1:]
                if (run.text[i] in special_chars) \
                        and (prev_char == ' '):
                    run.text = run.text[:i - 1] + run.text[i:]
                    limit = len(run.text)
        for pattern in incorrect_patterns:
            run.text = run.text.replace(pattern[0], pattern[1])
        if debug:
            print(run.text)
    if debug:
        print(paragraph.text)
    return paragraph


in_debug = False
if in_debug:
    document = Document("D:\\test.docx")
    length = len(document.paragraphs)
    for i in range(0, length):
        document.paragraphs[i] = remove_patterns(document.paragraphs[i], in_debug)
        loading(i, length, in_debug)
    now = datetime.datetime.now().strftime('%Y_%m_%d_%H_%M_%S')
    document.save(f'D:\\{now}.docx')
else:
    for file in os.listdir("."):
        if file.endswith(".docx"):
            document = Document(file)
            length = len(document.paragraphs)
            for i in range(0, length):
                document.paragraphs[i] = remove_patterns(document.paragraphs[i], in_debug)
                print(f'Formatting file: {file}')
                loading(i, length, in_debug)
            if not os.path.isdir("ToReview"):
                os.mkdir("ToReview")
            document.save(os.path.join("ToReview", file))
            if not os.path.isdir("Untouched"):
                os.mkdir("Untouched")
            shutil.move(file, os.path.join("Untouched", file))
